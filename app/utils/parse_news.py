from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from PIL import Image 
import io

class Parse_News():

    # Initializing article variables
    title = "This is article title"
    description = "this is news parser"
    highlight = False
    headerImg = "header Url"
    headerImageTitle = "header image title"

    # Dummy author
    author = {
                "displayName": "This is display name",
                "email": "this is email",
                "department": "this is department",
                "jobtitle": "this is job title"
            }

    # Empty paragraph dictionary
    paragraphs = {}

    # Node for images
    articleImage = {
        "filename": None,
        "imgTitle": None,
        "heigth": None,
        "width": None 
    }

    #Node for links
    articleLinks = {
        "link": None,
        "link_keyword": None
    }

    #Initializing the document file
    def initialize_doc(self, file):
        source_stream = io.BytesIO(file.read())
        document = Document(source_stream)
        document.save(source_stream)
        print(document)
        return document

    # Setting the values for author and description
    def setValues(self, author, description):
        self.author = author
        self.description = description
        print ("set values", self.author, self.description)

    # Creating the article dictionary for database parsing
    def createArticle(self):
        article = {
            "title": self.title,
            "author": {
                "displayName": self.author["displayName"],
                "email": self.author["email"],
                "department": self.author["department"],
                "jobtitle": self.author["jobtitle"]
            },
            "highlight": self.highlight,
            "description": self.description,
            "headerImgUrl": self.headerImg,
            "headerImgTitle": self.headerImageTitle,
            "paragraphs": self.paragraphs,
        }
        return article


    def parse_news(self, doc):
        print("parsing news")

        # Initializing indexes and 
        i = 0
        image_index = 0
        # Creating an array for images
        image_paragraphs = []
        rel_links = []

        # First we need a title so we have a fallback plan for images
        for para in doc.paragraphs:
            if para.style.name == "Heading 1":
                self.title = para.text

        # looping through images inside the document and saving them to image array
        rels = doc.part.rels
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                rel_links.append({"link": rel, "url": rels[rel]._target})
        # Checking for images
        for image in doc.inline_shapes:
            # First creating an articleImage object
            self.articleImage["filename"] = image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name
            self.articleImage["imgTitle"] = "" 
            self.articleImage["heigth"] = image.height.cm
            self.articleImage["width"] = image.width.cm
            if self.articleImage["filename"].startswith("Picture") and image_index == 0:
                self.articleImage["filename"] = self.title+".jpg"
                image_index += 1
            elif self.articleImage["filename"].startswith("Picture") and image_index > 0:
                self.articleImage["filename"] = self.title+image_index+".jpg"
                image_index += 1

            image_paragraphs.append(self.articleImage)

        # Looping through paragraphs inside the document
        for para in doc.paragraphs:
            links = {}
            link_index = 1
            # Checking for links
            link = None
            link_keyword = None
            hyperlink = para._p.xpath("./w:hyperlink")
            # Looping through links in paragraph and getting their row id's
            for hylink in hyperlink:
                hyperlink = hyperlink[0]
                hyperlink_rel_id = hyperlink.get(qn("r:id"))
                # Looping through the links parsed earlier
                for hlink in rel_links:
                    # If they match adding the link data to paragraph
                    print(hlink, hyperlink_rel_id)
                    if hlink["link"] == hyperlink_rel_id:
                        print("Adding a link")
                        links[link_index ] = {
                                "link": hlink["url"],
                                "link_keyword": hlink["url"]
                        }
                        link_index += 1
            # If links is still empty making it None for readability and null checking purposes  
            if links == {}:
                links = None
            # If paragraph is the main header
            if para.style.name == "Heading 1":
                self.title = para.text
            # If paragraph is the header image
            elif para.style.name == "Kuva" and i == 0 and len(image_paragraphs) > 0:
                self.headerImg = image_paragraphs[0]
                image_paragraphs.pop(0)
            # And if its neither we go to normal paragraph creation
            else:
                paraimg = None
                if 'graphicData' in para._p.xml and len(image_paragraphs) > 0:
                    paraimg = image_paragraphs[0]
                    paraimg["imgTitle"] = para.text
                    image_paragraphs.pop(0)

                newpara = {
                    "text": para.text,
                    "image": paraimg,
                    "style": para.style.name,
                    "links": links
                }
                #adding a paragraph
                i += 1
                self.paragraphs[i] = newpara    




